import streamlit as st
from docx import Document
from io import BytesIO
from stcn import get_dynamic_content  # 请确保正确导入你的爬虫函数

# 设置页面
st.set_page_config(page_icon="📰", page_title="证券时报News")
st.title('📰证券时报News🤖')
st.write("关键词：银行 | 钢铁 | 煤炭 | 能源 | 房地产 | 汽车 | 保险 | 金融")

key_words = ['银行', '钢铁', '煤炭', '能源', '房地产', '汽车', '保险', '金融']
news_all = []

def remove_duplicates(news_all):
    seen_titles = set()
    unique_news = []
    for news in news_all:
        title = news.get('title')
        if title not in seen_titles:
            unique_news.append(news)
            seen_titles.add(title)
    return unique_news

def filter_by_content_length(news_all):
    return [news for news in news_all if len(news.get('content', '')) >= 400]

# 按钮触发爬取
if st.button("开始爬取"):
    with st.spinner("正在爬取新闻..."):
        for key_word in key_words:
            news_data = get_dynamic_content(key_word)
            news_all.extend(news_data)

        news_all = remove_duplicates(news_all)
        news_all = filter_by_content_length(news_all)

    # 爬取完成后的状态更新
    if news_all:
        # 构建 Word 文档
        doc = Document()
        for key_word in key_words:
            doc.add_heading(key_word, level=1)
            for article in news_all:
                if article['keyword'] == key_word:
                    doc.add_heading(article.get('title', 'N/A'), level=2)
                    doc.add_paragraph(f"链接: {article.get('link', 'N/A')}")
                    doc.add_paragraph(f"日期: {article.get('date', 'N/A')}")
                    doc.add_paragraph(f"标签: {article.get('tags', 'N/A')}")
                    doc.add_paragraph(f"摘要: {article.get('summary', 'N/A')}")
                    doc.add_paragraph(f"内容: {article.get('content', 'N/A')}")
                    doc.add_paragraph("\n")

        # 保存Word文档到内存缓冲区
        news_word = BytesIO()
        doc.save(news_word)
        news_word.seek(0)

        # 提供下载链接
        st.success("新闻爬取完成！")
        st.download_button(label="下载新闻文档", data=news_word, file_name="news.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("没有爬取到任何新闻数据。")
